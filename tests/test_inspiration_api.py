import asyncio
import io
import os
import tempfile
import unittest
from pathlib import Path

from fastapi import FastAPI
from fastapi.testclient import TestClient
from sqlalchemy import create_engine
from sqlalchemy.orm import sessionmaker
from sqlalchemy.pool import StaticPool

from database import Base, get_db
from models import AIInterfaceSetting, Product, SellingPoint


class InspirationApiTests(unittest.TestCase):
    def setUp(self):
        from routers import inspiration

        self.inspiration = inspiration
        self.original_client = inspiration.ai_service.client
        self.original_chat = inspiration.ai_service.chat
        self.original_model = inspiration.ai_service.model
        self.original_clients = dict(getattr(inspiration.ai_service, "_clients", {}))
        self.original_provider_env = {
            key: os.environ.get(key)
            for key in ("ARK_API_KEY", "ARK_BASE_URL", "ARK_MODEL", "DEEPSEEK_API_KEY")
        }
        self.original_seedance_prompt_generator = getattr(inspiration, "seedance_prompt_generator", None)
        self.temp_documents = tempfile.TemporaryDirectory()
        self.original_document_dir = None
        if hasattr(inspiration, "inspiration_documents"):
            self.original_document_dir = getattr(inspiration.inspiration_documents, "DOCUMENT_OUTPUT_DIR", None)
            inspiration.inspiration_documents.DOCUMENT_OUTPUT_DIR = Path(self.temp_documents.name)
        self.had_ai_timeout = hasattr(inspiration, "INSPIRATION_AI_TIMEOUT_SECONDS")
        self.original_ai_timeout = getattr(inspiration, "INSPIRATION_AI_TIMEOUT_SECONDS", None)
        self.had_thinking_ai_timeout = hasattr(inspiration, "INSPIRATION_THINKING_AI_TIMEOUT_SECONDS")
        self.original_thinking_ai_timeout = getattr(inspiration, "INSPIRATION_THINKING_AI_TIMEOUT_SECONDS", None)
        self.engine = create_engine(
            "sqlite:///:memory:",
            connect_args={"check_same_thread": False},
            poolclass=StaticPool,
        )
        Base.metadata.create_all(bind=self.engine)
        self.Session = sessionmaker(bind=self.engine)
        self.db = self.Session()

        app = FastAPI()

        def override_db():
            try:
                yield self.db
            finally:
                pass

        app.dependency_overrides[get_db] = override_db
        app.include_router(inspiration.router, prefix="/api/inspiration")
        self.client = TestClient(app)

    def tearDown(self):
        self.inspiration.ai_service.client = self.original_client
        self.inspiration.ai_service.chat = self.original_chat
        self.inspiration.ai_service.model = self.original_model
        if hasattr(self.inspiration.ai_service, "_clients"):
            self.inspiration.ai_service._clients = dict(self.original_clients)
        for key, value in self.original_provider_env.items():
            if value is None:
                os.environ.pop(key, None)
            else:
                os.environ[key] = value
        if self.original_seedance_prompt_generator is None:
            if hasattr(self.inspiration, "seedance_prompt_generator"):
                delattr(self.inspiration, "seedance_prompt_generator")
        else:
            self.inspiration.seedance_prompt_generator = self.original_seedance_prompt_generator
        if self.original_document_dir is not None and hasattr(self.inspiration, "inspiration_documents"):
            self.inspiration.inspiration_documents.DOCUMENT_OUTPUT_DIR = self.original_document_dir
        self.temp_documents.cleanup()
        if self.had_ai_timeout:
            self.inspiration.INSPIRATION_AI_TIMEOUT_SECONDS = self.original_ai_timeout
        elif hasattr(self.inspiration, "INSPIRATION_AI_TIMEOUT_SECONDS"):
            delattr(self.inspiration, "INSPIRATION_AI_TIMEOUT_SECONDS")
        if self.had_thinking_ai_timeout:
            self.inspiration.INSPIRATION_THINKING_AI_TIMEOUT_SECONDS = self.original_thinking_ai_timeout
        elif hasattr(self.inspiration, "INSPIRATION_THINKING_AI_TIMEOUT_SECONDS"):
            delattr(self.inspiration, "INSPIRATION_THINKING_AI_TIMEOUT_SECONDS")
        self.db.close()
        Base.metadata.drop_all(bind=self.engine)
        self.engine.dispose()

    def _add_product(self, name, category, price, description, point):
        product = Product(
            name=name,
            category=category,
            price=price,
            brand="法采",
            description=description,
            status="active",
        )
        self.db.add(product)
        self.db.flush()
        self.db.add(SellingPoint(
            product_id=product.id,
            point_type="核心卖点",
            content=point,
            priority=1,
        ))
        self.db.commit()
        return product

    def _clear_provider_runtime(self):
        for key in ("ARK_API_KEY", "ARK_BASE_URL", "ARK_MODEL", "DEEPSEEK_API_KEY"):
            os.environ.pop(key, None)
        self.inspiration.ai_service.client = None
        self.inspiration.ai_service._clients = {}

    def test_chat_uses_interface_availability_when_only_ark_key_is_configured(self):
        original_env = {
            "ARK_API_KEY": os.environ.get("ARK_API_KEY"),
            "ARK_BASE_URL": os.environ.get("ARK_BASE_URL"),
            "ARK_MODEL": os.environ.get("ARK_MODEL"),
            "DEEPSEEK_API_KEY": os.environ.get("DEEPSEEK_API_KEY"),
        }
        os.environ["ARK_API_KEY"] = "ark-only-secret"
        os.environ["ARK_BASE_URL"] = "https://ark.example.test/api/v3"
        os.environ["ARK_MODEL"] = "ep-unit-test"
        os.environ.pop("DEEPSEEK_API_KEY", None)
        self.inspiration.ai_service.client = None
        self.inspiration.ai_service._clients = {}

        async def fake_chat(*args, **kwargs):
            self.assertEqual(kwargs.get("interface_key"), "inspiration_chat")
            self.assertIsNone(kwargs.get("model"))
            return {"content": "火山方舟正常回复", "reasoning": "", "model": "ep-unit-test"}

        self.inspiration.ai_service.chat = fake_chat
        try:
            response = self.client.post(
                "/api/inspiration/chat",
                json={"message": "帮我生成一个活动选题", "tool_mode": "chat"},
            )
        finally:
            for key, value in original_env.items():
                if value is None:
                    os.environ.pop(key, None)
                else:
                    os.environ[key] = value

        self.assertEqual(response.status_code, 200)
        data = response.json()
        self.assertEqual(data["mode"], "ai")
        self.assertEqual(data["answer"], "火山方舟正常回复")
        self.assertEqual(data["model"], "ep-unit-test")

    def test_chat_requires_non_empty_message(self):
        response = self.client.post("/api/inspiration/chat", json={"message": "   "})

        self.assertEqual(response.status_code, 422)

    def test_chat_returns_ai_answer_when_service_responds(self):
        self.inspiration.ai_service.client = object()
        self.inspiration.ai_service.model = "deepseek-chat"

        async def fake_chat(messages, temperature=0.7, allow_fallback=False, model=None, thinking=False, return_reasoning=False, **kwargs):
            self.assertFalse(allow_fallback)
            self.assertIsNone(model)
            self.assertFalse(thinking)
            self.assertTrue(return_reasoning)
            self.assertEqual(messages[0]["role"], "system")
            self.assertIn("法采新媒体运营AI工作助手", messages[0]["content"])
            self.assertEqual(messages[-1], {"role": "user", "content": "帮我想 3 个新品短视频开头"})
            return {"content": "这里是 3 个开头。", "reasoning": "", "model": "fake-interface-model"}

        self.inspiration.ai_service.chat = fake_chat

        response = self.client.post(
            "/api/inspiration/chat",
            json={"message": "帮我想 3 个新品短视频开头"},
        )

        self.assertEqual(response.status_code, 200)
        data = response.json()
        self.assertEqual(data["answer"], "这里是 3 个开头。")
        self.assertEqual(data["mode"], "ai")
        self.assertEqual(data["model"], "fake-interface-model")
        self.assertFalse(data["product_context_used"])
        self.assertEqual(data["products"], [])
        self.assertEqual(data["tool_mode"], "chat")
        self.assertEqual(data["reasoning"], "")
        self.assertEqual(data["sources"], [])
        self.assertEqual(data["attachments_used"], [])

    def test_chat_thinking_mode_uses_interface_default_model_and_returns_reasoning(self):
        self.inspiration.ai_service.client = object()
        captured = {}

        async def fake_chat(messages, temperature=0.7, allow_fallback=False, model=None, thinking=False, reasoning_effort=None, return_reasoning=False, **kwargs):
            captured["model"] = model
            captured["thinking"] = thinking
            captured["reasoning_effort"] = reasoning_effort
            captured["return_reasoning"] = return_reasoning
            captured["messages"] = messages
            return {"content": "这是思考后的方案。", "reasoning": "先拆目标，再比较打法。", "model": "fake-interface-model"}

        self.inspiration.ai_service.chat = fake_chat

        response = self.client.post(
            "/api/inspiration/chat",
            json={"message": "帮我认真分析这个活动怎么做", "tool_mode": "thinking"},
        )

        self.assertEqual(response.status_code, 200)
        data = response.json()
        self.assertIsNone(captured["model"])
        self.assertTrue(captured["thinking"])
        self.assertEqual(captured["reasoning_effort"], "high")
        self.assertTrue(captured["return_reasoning"])
        self.assertIn("思考模式", captured["messages"][0]["content"])
        self.assertEqual(data["answer"], "这是思考后的方案。")
        self.assertEqual(data["model"], "fake-interface-model")
        self.assertEqual(data["tool_mode"], "thinking")
        self.assertEqual(data["reasoning"], "先拆目标，再比较打法。")

    def test_model_status_question_answers_locally_without_ai_or_product_context(self):
        self.db.add(AIInterfaceSetting(
            interface_key="inspiration_tools",
            provider="doubao",
            model="doubao-2.1-pro",
            max_tokens=3600,
        ))
        self._add_product(
            "2元盒装",
            "烘焙配件",
            2.03,
            "一次性盒装配件",
            "适合低价促销组合。",
        )
        self.inspiration.ai_service.client = object()

        async def fail_if_called(*args, **kwargs):
            raise AssertionError("model status questions should not call AI")

        self.inspiration.ai_service.chat = fail_if_called

        response = self.client.post(
            "/api/inspiration/chat",
            json={"message": "你用的是豆包2.1还是deepseek", "tool_mode": "thinking"},
        )

        self.assertEqual(response.status_code, 200)
        data = response.json()
        self.assertEqual(data["mode"], "ai")
        self.assertEqual(data["model"], "doubao-2.1-pro")
        self.assertEqual(data["tool_mode"], "thinking")
        self.assertFalse(data["product_context_used"])
        self.assertEqual(data["products"], [])
        self.assertIn("当前功能：思考模式", data["answer"])
        self.assertIn("服务商：豆包 / 火山方舟", data["answer"])
        self.assertIn("模型：doubao-2.1-pro", data["answer"])
        self.assertIn("不是 DeepSeek", data["answer"])

    def test_chat_returns_fallback_when_ai_call_times_out(self):
        self.inspiration.ai_service.client = object()
        self.inspiration.INSPIRATION_AI_TIMEOUT_SECONDS = 0.01

        async def slow_chat(messages, temperature=0.7, allow_fallback=False, **kwargs):
            await asyncio.sleep(0.05)
            return {"content": "late answer", "reasoning": "", "model": "slow-model"}

        self.inspiration.ai_service.chat = slow_chat

        response = self.client.post(
            "/api/inspiration/chat",
            json={"message": "timeout test"},
        )

        self.assertEqual(response.status_code, 200)
        data = response.json()
        self.assertEqual(data["mode"], "fallback")
        self.assertNotEqual(data["answer"], "late answer")
        self.assertEqual(data["tool_mode"], "chat")
        self.assertIn("AI 响应超时", data["answer"])

    def test_thinking_mode_uses_longer_ai_timeout(self):
        self.inspiration.ai_service.client = object()
        self.inspiration.INSPIRATION_AI_TIMEOUT_SECONDS = 0.01
        self.inspiration.INSPIRATION_THINKING_AI_TIMEOUT_SECONDS = 0.2

        async def slow_but_valid_chat(messages, temperature=0.7, allow_fallback=False, **kwargs):
            await asyncio.sleep(0.05)
            return {"content": "thinking answer", "reasoning": "reasoning trace", "model": "deepseek-v4-pro"}

        self.inspiration.ai_service.chat = slow_but_valid_chat

        response = self.client.post(
            "/api/inspiration/chat",
            json={"message": "thinking timeout test", "tool_mode": "thinking"},
        )

        self.assertEqual(response.status_code, 200)
        data = response.json()
        self.assertEqual(data["mode"], "ai")
        self.assertEqual(data["answer"], "thinking answer")
        self.assertEqual(data["tool_mode"], "thinking")

    def test_thinking_mode_passes_extended_request_timeout_to_ai_service(self):
        self.inspiration.ai_service.client = object()
        self.inspiration.INSPIRATION_THINKING_AI_TIMEOUT_SECONDS = 240.0
        captured = {}

        async def fake_chat(messages, temperature=0.7, allow_fallback=False, **kwargs):
            captured["request_timeout"] = kwargs.get("request_timeout")
            return {"content": "thinking answer", "reasoning": "", "model": "deepseek-v4-pro"}

        self.inspiration.ai_service.chat = fake_chat

        response = self.client.post(
            "/api/inspiration/chat",
            json={"message": "thinking timeout test", "tool_mode": "thinking"},
        )

        self.assertEqual(response.status_code, 200)
        self.assertEqual(response.json()["mode"], "ai")
        self.assertEqual(captured["request_timeout"], 240.0)

    def test_research_mode_adds_web_sources_to_prompt(self):
        self.inspiration.ai_service.client = object()
        captured = {}
        original_search = getattr(self.inspiration, "search_web", None)

        async def fake_search(query, max_results=5):
            self.assertIn("烘焙", query)
            return [
                {"title": "烘焙内容趋势", "url": "https://example.com/trend", "snippet": "用户更关注低成本出片。"},
                {"title": "直播运营方法", "url": "https://example.com/live", "snippet": "直播间促单需要明确价格锚点。"},
            ]

        async def fake_chat(messages, temperature=0.7, allow_fallback=False, model=None, thinking=False, return_reasoning=False, **kwargs):
            captured["messages"] = messages
            captured["model"] = model
            captured["thinking"] = thinking
            return {"content": "结合外网资料的研究结论。", "reasoning": "", "model": "fake-interface-model"}

        self.inspiration.search_web = fake_search
        self.inspiration.ai_service.chat = fake_chat
        try:
            response = self.client.post(
                "/api/inspiration/chat",
                json={"message": "研究一下烘焙短视频趋势", "tool_mode": "research"},
            )
        finally:
            if original_search is None:
                delattr(self.inspiration, "search_web")
            else:
                self.inspiration.search_web = original_search

        self.assertEqual(response.status_code, 200)
        data = response.json()
        self.assertIsNone(captured["model"])
        self.assertFalse(captured["thinking"])
        self.assertIn("外网搜索结果", captured["messages"][-1]["content"])
        self.assertIn("烘焙内容趋势", captured["messages"][-1]["content"])
        self.assertEqual(data["tool_mode"], "research")
        self.assertEqual(len(data["sources"]), 2)
        self.assertEqual(data["sources"][0]["url"], "https://example.com/trend")

    def test_analysis_mode_uses_attachments_and_web_sources(self):
        self.inspiration.ai_service.client = object()
        captured = {}
        original_search = getattr(self.inspiration, "search_web", None)

        async def fake_search(query, max_results=5):
            return [{"title": "行业均值", "url": "https://example.com/data", "snippet": "短视频互动率均值约 3%。"}]

        async def fake_chat(messages, temperature=0.7, allow_fallback=False, model=None, thinking=False, return_reasoning=False, **kwargs):
            captured["messages"] = messages
            return {"content": "数据分析结论。", "reasoning": "", "model": model}

        self.inspiration.search_web = fake_search
        self.inspiration.ai_service.chat = fake_chat
        try:
            response = self.client.post(
                "/api/inspiration/chat",
                json={
                    "message": "分析这份投放数据",
                    "tool_mode": "analysis",
                    "attachments": [
                        {
                            "filename": "投放数据.csv",
                            "file_type": "csv",
                            "text": "日期,播放量,成交\n2026-06-01,1000,20",
                            "char_count": 25,
                        }
                    ],
                },
            )
        finally:
            if original_search is None:
                delattr(self.inspiration, "search_web")
            else:
                self.inspiration.search_web = original_search

        self.assertEqual(response.status_code, 200)
        data = response.json()
        self.assertIn("附件资料", captured["messages"][-1]["content"])
        self.assertIn("投放数据.csv", captured["messages"][-1]["content"])
        self.assertIn("外网搜索结果", captured["messages"][-1]["content"])
        self.assertEqual(data["tool_mode"], "analysis")
        self.assertEqual(data["attachments_used"][0]["filename"], "投放数据.csv")
        self.assertEqual(data["sources"][0]["title"], "行业均值")

    def test_seedance_mode_uses_message_as_script_without_attachments(self):
        class FakeSeedanceGenerator:
            def __init__(self):
                self.calls = []

            async def generate(self, *, script_content, requirements=None, db=None):
                self.calls.append({
                    "script_content": script_content,
                    "requirements": requirements,
                    "db": db,
                })
                return {
                    "prompt_text": "画面1：竖屏9:16，真人口播展示糖珠，无字幕无水印。",
                    "items": [],
                    "source": "ai",
                }

        fake = FakeSeedanceGenerator()
        self.inspiration.seedance_prompt_generator = fake
        self.inspiration.ai_service.client = None

        response = self.client.post(
            "/api/inspiration/chat",
            json={"message": "（口播）介绍这款糖珠不用挑尺寸。", "tool_mode": "seedance"},
        )

        self.assertEqual(response.status_code, 200)
        data = response.json()
        self.assertEqual(data["mode"], "ai")
        self.assertEqual(data["tool_mode"], "seedance")
        self.assertEqual(data["answer"], "画面1：竖屏9:16，真人口播展示糖珠，无字幕无水印。")
        self.assertEqual(fake.calls[0]["script_content"], "（口播）介绍这款糖珠不用挑尺寸。")
        self.assertIsNone(fake.calls[0]["requirements"])
        self.assertIs(fake.calls[0]["db"], self.db)

    def test_seedance_mode_uses_attachment_text_as_script_and_message_as_requirements(self):
        class FakeSeedanceGenerator:
            def __init__(self):
                self.calls = []

            async def generate(self, *, script_content, requirements=None, db=None):
                self.calls.append({
                    "script_content": script_content,
                    "requirements": requirements,
                    "db": db,
                })
                return {
                    "prompt_text": "画面1：竖屏9:16，展示新品包装，无字幕无水印。",
                    "items": [],
                    "source": "ai",
                }

        fake = FakeSeedanceGenerator()
        self.inspiration.seedance_prompt_generator = fake

        response = self.client.post(
            "/api/inspiration/chat",
            json={
                "message": "生成时突出门店真实口播",
                "tool_mode": "seedance",
                "attachments": [
                    {
                        "filename": "脚本.txt",
                        "file_type": "txt",
                        "text": "（主播出镜）这款糖珠拆开就能用。",
                        "char_count": 18,
                    }
                ],
            },
        )

        self.assertEqual(response.status_code, 200)
        data = response.json()
        self.assertEqual(data["mode"], "ai")
        self.assertEqual(data["tool_mode"], "seedance")
        self.assertEqual(data["attachments_used"][0]["filename"], "脚本.txt")
        self.assertEqual(fake.calls[0]["script_content"], "（主播出镜）这款糖珠拆开就能用。")
        self.assertEqual(fake.calls[0]["requirements"], "生成时突出门店真实口播")

    def test_seedance_mode_returns_generation_error_without_chat_fallback(self):
        from services.seedance_prompt_generator import SeedancePromptGenerationError

        class FailingSeedanceGenerator:
            async def generate(self, *, script_content, requirements=None, db=None):
                raise SeedancePromptGenerationError(503, "DeepSeek V4 Pro / 脚本改写配置不可用")

        async def fail_if_called(*args, **kwargs):
            raise AssertionError("Seedance mode should not call regular inspiration chat")

        self.inspiration.seedance_prompt_generator = FailingSeedanceGenerator()
        self.inspiration.ai_service.chat = fail_if_called

        response = self.client.post(
            "/api/inspiration/chat",
            json={"message": "（口播）介绍糖珠", "tool_mode": "seedance"},
        )

        self.assertEqual(response.status_code, 503)
        self.assertIn("DeepSeek V4 Pro", response.json()["detail"])

    def test_attachment_upload_extracts_text_file(self):
        response = self.client.post(
            "/api/inspiration/attachments",
            files={"file": ("brief.txt", b"\xe7\x81\xb5\xe6\x84\x9f\xe9\x99\x84\xe4\xbb\xb6", "text/plain")},
        )

        self.assertEqual(response.status_code, 200)
        data = response.json()
        self.assertEqual(data["filename"], "brief.txt")
        self.assertEqual(data["file_type"], "txt")
        self.assertIn("灵感附件", data["text"])

    def test_attachment_upload_extracts_word_docx(self):
        from docx import Document

        document = Document()
        document.add_paragraph("Word 附件内容")
        handle = io.BytesIO()
        document.save(handle)
        handle.seek(0)

        response = self.client.post(
            "/api/inspiration/attachments",
            files={
                "file": (
                    "brief.docx",
                    handle.getvalue(),
                    "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                )
            },
        )

        self.assertEqual(response.status_code, 200)
        data = response.json()
        self.assertEqual(data["file_type"], "docx")
        self.assertIn("Word 附件内容", data["text"])

    def test_attachment_upload_rejects_images_for_now(self):
        response = self.client.post(
            "/api/inspiration/attachments",
            files={"file": ("photo.png", b"not-real-image", "image/png")},
        )

        self.assertEqual(response.status_code, 415)
        self.assertIn("图片", response.json()["detail"])

    def test_attachment_upload_rejects_oversize_before_extracting(self):
        original_limit = getattr(self.inspiration, "MAX_ATTACHMENT_BYTES", None)
        original_extract = self.inspiration.extract_attachment_text
        self.inspiration.MAX_ATTACHMENT_BYTES = 4

        def fail_if_called(*args, **kwargs):
            raise AssertionError("oversize uploads should be rejected before extraction")

        self.inspiration.extract_attachment_text = fail_if_called
        try:
            response = self.client.post(
                "/api/inspiration/attachments",
                files={"file": ("too-large.txt", b"12345", "text/plain")},
            )
        finally:
            self.inspiration.extract_attachment_text = original_extract
            if original_limit is None:
                delattr(self.inspiration, "MAX_ATTACHMENT_BYTES")
            else:
                self.inspiration.MAX_ATTACHMENT_BYTES = original_limit

        self.assertEqual(response.status_code, 413)

    def test_document_export_creates_downloadable_word_file(self):
        self._clear_provider_runtime()

        response = self.client.post(
            "/api/inspiration/documents",
            json={
                "message": "帮我整理一个奶冻粉活动方案",
                "answer": "奶冻粉活动方案\n\n1. 主打低成本出片。\n2. 直播间强调稳定成型。",
                "history": [
                    {"role": "user", "content": "帮我整理一个奶冻粉活动方案"},
                    {"role": "assistant", "content": "奶冻粉活动方案"},
                ],
                "attachments": [
                    {
                        "filename": "活动节奏.docx",
                        "file_type": "docx",
                        "text": "第一周预热，第二周直播转化。",
                        "char_count": 16,
                    }
                ],
                "products": [
                    {
                        "product_id": 12,
                        "name": "奶冻粉",
                        "category": "烘焙原料",
                        "price": 18.8,
                    }
                ],
            },
        )

        self.assertEqual(response.status_code, 200)
        data = response.json()
        self.assertTrue(data["filename"].endswith(".docx"))
        self.assertTrue(data["download_url"].startswith("/api/inspiration/documents/"))
        self.assertIn("奶冻粉", data["title"])

        download = self.client.get(data["download_url"])
        self.assertEqual(download.status_code, 200)
        self.assertIn(
            "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            download.headers["content-type"],
        )

        from docx import Document

        document = Document(io.BytesIO(download.content))
        text = "\n".join(paragraph.text for paragraph in document.paragraphs)
        table_text = "\n".join(cell.text for table in document.tables for row in table.rows for cell in row.cells)
        combined = text + "\n" + table_text
        self.assertIn("奶冻粉活动方案", combined)
        self.assertIn("参考产品", combined)
        self.assertIn("奶冻粉", combined)
        self.assertIn("活动节奏.docx", combined)

    def test_document_export_accepts_long_answer_history(self):
        self._clear_provider_runtime()
        long_answer = "战役流程优化建议\n" + ("监控计划执行、数据表现、风险状态和资源占用。\n" * 180)

        response = self.client.post(
            "/api/inspiration/documents",
            json={
                "message": "根据智能纪要优化战役运作流程",
                "answer": long_answer,
                "history": [
                    {"role": "user", "content": "根据智能纪要优化战役运作流程"},
                    {"role": "assistant", "content": long_answer},
                ],
            },
        )

        self.assertEqual(response.status_code, 200)
        data = response.json()
        self.assertTrue(data["download_url"].endswith("/download"))

    def test_document_download_rejects_unsafe_filename(self):
        response = self.client.get("/api/inspiration/documents/..%2Fsecret.docx/download")

        self.assertEqual(response.status_code, 404)

    def test_chat_uses_product_context_when_message_mentions_product_intent(self):
        product = self._add_product(
            "水性色素",
            "烘焙调色",
            18.59,
            "适合蛋糕调色和翻糖上色",
            "少量即可上色，适合烘焙门店做调色备货。",
        )
        self.inspiration.ai_service.client = object()
        captured = {}

        async def fake_chat(messages, temperature=0.7, allow_fallback=False, **kwargs):
            captured["messages"] = messages
            return "可以围绕水性色素做一个调色前后对比脚本。"

        self.inspiration.ai_service.chat = fake_chat

        response = self.client.post(
            "/api/inspiration/chat",
            json={"message": "帮我想一个调色产品短视频脚本"},
        )

        self.assertEqual(response.status_code, 200)
        data = response.json()
        self.assertTrue(data["product_context_used"])
        self.assertEqual(data["products"][0]["product_id"], product.id)
        self.assertEqual(data["products"][0]["name"], "水性色素")
        self.assertIn("产品资料", captured["messages"][-1]["content"])
        self.assertIn("水性色素", captured["messages"][-1]["content"])

    def test_chat_product_context_mode_always_forces_product_context_lookup(self):
        self.inspiration.ai_service.client = object()
        captured = {}
        original_finder = self.inspiration.find_product_context_for_inspiration

        def fake_product_context(query, db, *, limit=6, force=False):
            captured["query"] = query
            captured["limit"] = limit
            captured["force"] = force
            return {
                "used": bool(force),
                "context": "1. 产品：奶冻粉\n品类：烘焙夹心\n卖点：口感稳定，适合门店做奶冻夹心内容。",
                "products": [
                    {
                        "product_id": 88,
                        "name": "奶冻粉",
                        "category": "烘焙夹心",
                        "price": 12.71,
                    }
                ] if force else [],
            }

        async def fake_chat(messages, temperature=0.7, allow_fallback=False, **kwargs):
            captured["prompt"] = messages[-1]["content"]
            return "已基于奶冻粉资料给出内容建议。"

        self.inspiration.find_product_context_for_inspiration = fake_product_context
        self.inspiration.ai_service.chat = fake_chat
        try:
            response = self.client.post(
                "/api/inspiration/chat",
                json={
                    "message": "今天拍什么内容",
                    "product_context_mode": "always",
                },
            )
        finally:
            self.inspiration.find_product_context_for_inspiration = original_finder

        self.assertEqual(response.status_code, 200)
        data = response.json()
        self.assertTrue(captured["force"])
        self.assertEqual(captured["limit"], 6)
        self.assertIn("产品资料", captured["prompt"])
        self.assertTrue(data["product_context_used"])
        self.assertEqual(data["products"][0]["name"], "奶冻粉")

    def test_chat_uses_product_context_for_product_names_and_selling_point_keywords(self):
        fondant = self._add_product(
            "翻糖压片",
            "烘焙装饰",
            31.53,
            "彩色翻糖片，可用于造型蛋糕装饰。",
            "适合做蛋糕表面造型和节日款装饰。",
        )
        puree = self._add_product(
            "夹心果泥",
            "烘焙夹心",
            49.29,
            "果肉含量高，适合蛋糕夹心和奶油调味。",
            "开袋即用，适合活动款蛋糕做夹心卖点。",
        )
        self.inspiration.ai_service.client = object()

        async def fake_chat(messages, temperature=0.7, allow_fallback=False, **kwargs):
            return "已结合产品资料给出内容方向。"

        self.inspiration.ai_service.chat = fake_chat

        fondant_response = self.client.post(
            "/api/inspiration/chat",
            json={"message": "翻糖怎么做内容选题"},
        )
        puree_response = self.client.post(
            "/api/inspiration/chat",
            json={"message": "果泥适合什么活动文案"},
        )

        self.assertEqual(fondant_response.status_code, 200)
        self.assertEqual(puree_response.status_code, 200)
        self.assertIn(fondant.id, [item["product_id"] for item in fondant_response.json()["products"]])
        self.assertIn(puree.id, [item["product_id"] for item in puree_response.json()["products"]])

    def test_chat_returns_local_fallback_when_ai_unavailable(self):
        self._clear_provider_runtime()

        response = self.client.post(
            "/api/inspiration/chat",
            json={"message": "今天拍什么内容？"},
        )

        self.assertEqual(response.status_code, 200)
        data = response.json()
        self.assertEqual(data["mode"], "fallback")
        self.assertIn("AI 服务暂时不可用", data["answer"])
        self.assertIn("今天拍什么内容？", data["answer"])

    def test_chat_ai_unavailable_still_returns_product_references(self):
        product = self._add_product(
            "水性色素",
            "烘焙调色",
            18.59,
            "适合蛋糕调色和翻糖上色",
            "少量即可上色，适合烘焙门店做调色备货。",
        )
        self._clear_provider_runtime()

        response = self.client.post(
            "/api/inspiration/chat",
            json={"message": "调色产品怎么拍"},
        )

        self.assertEqual(response.status_code, 200)
        data = response.json()
        self.assertEqual(data["mode"], "fallback")
        self.assertTrue(data["product_context_used"])
        self.assertEqual(data["products"][0]["product_id"], product.id)

    def test_chat_sends_only_recent_valid_history(self):
        self.inspiration.ai_service.client = object()
        captured = {}

        async def fake_chat(messages, temperature=0.7, allow_fallback=False, **kwargs):
            captured["messages"] = messages
            return "已结合上下文回答。"

        self.inspiration.ai_service.chat = fake_chat
        history = []
        for index in range(20):
            role = "user" if index % 2 == 0 else "assistant"
            history.append({"role": role, "content": f"历史 {index}"})
        history.append({"role": "system", "content": "should be ignored"})
        history.append({"role": "user", "content": ""})

        response = self.client.post(
            "/api/inspiration/chat",
            json={"message": "继续", "history": history},
        )

        self.assertEqual(response.status_code, 200)
        roles = [item["role"] for item in captured["messages"]]
        contents = [item["content"] for item in captured["messages"]]
        self.assertEqual(roles[0], "system")
        self.assertEqual(roles[-1], "user")
        self.assertEqual(contents[-1], "继续")
        self.assertNotIn("should be ignored", contents)
        self.assertLessEqual(len(captured["messages"]), 14)
        self.assertIn("历史 8", contents)
        self.assertIn("历史 19", contents)


if __name__ == "__main__":
    unittest.main()
