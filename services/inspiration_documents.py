"""Word document generation for AI Work chat answers."""
from __future__ import annotations

from dataclasses import dataclass
from datetime import datetime
from pathlib import Path
import re
from uuid import uuid4

from docx import Document
from docx.shared import Pt


DOCUMENT_OUTPUT_DIR = Path("data/generated_documents")
DOCX_MEDIA_TYPE = "application/vnd.openxmlformats-officedocument.wordprocessingml.document"


@dataclass
class GeneratedDocument:
    title: str
    filename: str
    path: Path


def _clean_text(value: object, limit: int = 20000) -> str:
    text = str(value or "").replace("\x00", "").strip()
    return text[:limit]


def _document_title(message: str, content: str, explicit_title: str = "") -> str:
    for candidate in (explicit_title, message, content.splitlines()[0] if content else ""):
        clean = re.sub(r"\s+", " ", str(candidate or "")).strip(" -_#：:")
        if clean:
            return clean[:42]
    return "AI工作文档"


def _safe_filename(title: str) -> str:
    slug = re.sub(r"[^\w\u4e00-\u9fff-]+", "-", title, flags=re.UNICODE).strip("-_")
    slug = slug[:36] or "ai-work-document"
    timestamp = datetime.now().strftime("%Y%m%d-%H%M%S")
    return f"{timestamp}-{slug}-{uuid4().hex[:8]}.docx"


def _set_default_font(document: Document) -> None:
    style = document.styles["Normal"]
    style.font.name = "Microsoft YaHei"
    style.font.size = Pt(10.5)


def _add_markdownish_content(document: Document, content: str) -> None:
    lines = _clean_text(content, 50000).splitlines()
    pending_blank = False
    for raw_line in lines:
        line = raw_line.strip()
        if not line:
            pending_blank = True
            continue
        if line.startswith("### "):
            document.add_heading(line[4:].strip(), level=3)
        elif line.startswith("## "):
            document.add_heading(line[3:].strip(), level=2)
        elif line.startswith("# "):
            document.add_heading(line[2:].strip(), level=1)
        elif re.match(r"^[-*]\s+", line):
            document.add_paragraph(re.sub(r"^[-*]\s+", "", line), style="List Bullet")
        elif re.match(r"^\d+[.)、]\s+", line):
            document.add_paragraph(re.sub(r"^\d+[.)、]\s+", "", line), style="List Number")
        else:
            if pending_blank and document.paragraphs:
                document.add_paragraph("")
            document.add_paragraph(line)
        pending_blank = False


def _format_price(value: object) -> str:
    try:
        number = float(value)
    except (TypeError, ValueError):
        return ""
    return f"¥{number:.2f}".rstrip("0").rstrip(".")


def _add_products(document: Document, products: list[dict]) -> None:
    rows = [product for product in products[:6] if product.get("name")]
    if not rows:
        return
    document.add_heading("参考产品", level=1)
    table = document.add_table(rows=1, cols=4)
    table.style = "Table Grid"
    headers = ["产品", "品类", "价格", "产品ID"]
    for index, label in enumerate(headers):
        table.rows[0].cells[index].text = label
    for product in rows:
        cells = table.add_row().cells
        cells[0].text = _clean_text(product.get("name"), 120)
        cells[1].text = _clean_text(product.get("category"), 120)
        cells[2].text = _format_price(product.get("price"))
        cells[3].text = _clean_text(product.get("product_id"), 40)


def _add_attachments(document: Document, attachments: list[dict]) -> None:
    rows = [attachment for attachment in attachments[:6] if attachment.get("filename")]
    if not rows:
        return
    document.add_heading("参考附件", level=1)
    for attachment in rows:
        file_type = _clean_text(attachment.get("file_type"), 20)
        filename = _clean_text(attachment.get("filename"), 160)
        label = f"{filename}（{file_type}）" if file_type else filename
        document.add_paragraph(label, style="List Bullet")


def create_inspiration_document(
    *,
    message: str,
    content: str,
    products: list[dict] | None = None,
    attachments: list[dict] | None = None,
    title: str = "",
) -> GeneratedDocument:
    body = _clean_text(content, 60000)
    if not body:
        raise ValueError("文档内容不能为空")
    document_title = _document_title(message, body, title)
    DOCUMENT_OUTPUT_DIR.mkdir(parents=True, exist_ok=True)
    filename = _safe_filename(document_title)
    path = DOCUMENT_OUTPUT_DIR / filename

    document = Document()
    _set_default_font(document)
    document.add_heading(document_title, level=0)
    document.add_paragraph(f"生成时间：{datetime.now().strftime('%Y-%m-%d %H:%M')}")
    if _clean_text(message, 4000):
        document.add_heading("用户需求", level=1)
        document.add_paragraph(_clean_text(message, 4000))
    document.add_heading("文档内容", level=1)
    _add_markdownish_content(document, body)
    _add_products(document, products or [])
    _add_attachments(document, attachments or [])
    document.save(path)
    return GeneratedDocument(title=document_title, filename=filename, path=path)


def resolve_document_path(filename: str) -> Path:
    safe_name = Path(filename or "").name
    if safe_name != filename or not safe_name.endswith(".docx"):
        raise ValueError("非法文件名")
    path = DOCUMENT_OUTPUT_DIR / safe_name
    if not path.exists() or not path.is_file():
        raise FileNotFoundError(filename)
    return path
