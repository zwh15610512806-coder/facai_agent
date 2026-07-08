"""Attachment text extraction for the Inspiration chat."""
from __future__ import annotations

import base64
from dataclasses import dataclass
from io import BytesIO
from pathlib import Path
import os
import tempfile
import uuid


MAX_ATTACHMENT_BYTES = 12 * 1024 * 1024
MAX_EXTRACTED_CHARS = 24000

TEXT_EXTENSIONS = {".txt", ".md", ".json", ".csv"}
SUPPORTED_EXTENSIONS = TEXT_EXTENSIONS | {".pdf", ".docx", ".xlsx"}
IMAGE_EXTENSIONS = {".jpg", ".jpeg", ".png", ".webp"}
ALL_IMAGE_EXTENSIONS = IMAGE_EXTENSIONS | {".gif", ".bmp"}
IMAGE_MIME_BY_EXTENSION = {
    ".jpg": "image/jpeg",
    ".jpeg": "image/jpeg",
    ".png": "image/png",
    ".webp": "image/webp",
}
IMAGE_EXTENSION_BY_MIME = {
    "image/jpeg": ".jpg",
    "image/png": ".png",
    "image/webp": ".webp",
}
INSPIRATION_ATTACHMENT_DIR = Path(tempfile.gettempdir()) / "facai-inspiration-attachments"


class AttachmentExtractionError(Exception):
    """Raised when an uploaded attachment cannot be converted into text."""

    def __init__(self, message: str, status_code: int = 400):
        super().__init__(message)
        self.status_code = status_code


@dataclass
class ExtractedAttachment:
    filename: str
    file_type: str
    text: str
    char_count: int
    kind: str = "text"
    attachment_id: str = ""
    mime_type: str = ""
    preview_url: str = ""


def _extension(filename: str) -> str:
    return Path(filename or "").suffix.lower()


def _trim_text(text: str) -> str:
    text = (text or "").replace("\x00", "").strip()
    return text[:MAX_EXTRACTED_CHARS]


def _image_mime_from_bytes(data: bytes) -> str:
    if data.startswith(b"\xff\xd8\xff"):
        return "image/jpeg"
    if data.startswith(b"\x89PNG\r\n\x1a\n"):
        return "image/png"
    if len(data) >= 12 and data[:4] == b"RIFF" and data[8:12] == b"WEBP":
        return "image/webp"
    return ""


def _normalized_content_type(content_type: str) -> str:
    return (content_type or "").split(";", 1)[0].strip().lower()


def _image_dir() -> Path:
    INSPIRATION_ATTACHMENT_DIR.mkdir(parents=True, exist_ok=True)
    return INSPIRATION_ATTACHMENT_DIR


def _safe_uuid(value: str) -> str:
    try:
        return str(uuid.UUID(str(value or "").strip()))
    except (TypeError, ValueError) as exc:
        raise AttachmentExtractionError("图片附件已失效，请重新上传。", 400) from exc


def _save_image_attachment(filename: str, content_type: str, data: bytes) -> ExtractedAttachment:
    if len(data) > MAX_ATTACHMENT_BYTES:
        raise AttachmentExtractionError("文件过大，请控制在 12MB 以内。", 413)
    if not data:
        raise AttachmentExtractionError("文件为空，请重新选择。", 400)

    ext = _extension(filename)
    declared_mime = _normalized_content_type(content_type)
    detected_mime = _image_mime_from_bytes(data)
    if ext in ALL_IMAGE_EXTENSIONS and ext not in IMAGE_EXTENSIONS:
        raise AttachmentExtractionError("暂不支持该图片格式，请上传 JPG、PNG 或 WebP 图片。", 415)
    if declared_mime.startswith("image/") and declared_mime not in IMAGE_EXTENSION_BY_MIME:
        raise AttachmentExtractionError("暂不支持该图片格式，请上传 JPG、PNG 或 WebP 图片。", 415)
    if not detected_mime:
        raise AttachmentExtractionError("图片文件解析失败，请重新选择图片。", 400)
    if ext and ext not in IMAGE_EXTENSIONS:
        raise AttachmentExtractionError("暂不支持该文件类型，请上传 PDF、Word、TXT、Markdown、JSON、CSV、XLSX 或图片。", 415)
    if ext and IMAGE_MIME_BY_EXTENSION.get(ext) != detected_mime:
        raise AttachmentExtractionError("图片扩展名和实际格式不一致，请重新选择图片。", 400)
    if declared_mime and declared_mime.startswith("image/") and declared_mime != detected_mime:
        raise AttachmentExtractionError("图片 MIME 类型和实际格式不一致，请重新选择图片。", 400)

    attachment_id = str(uuid.uuid4())
    storage_ext = IMAGE_EXTENSION_BY_MIME[detected_mime]
    safe_name = Path(filename or f"clipboard{storage_ext}").name
    if not _extension(safe_name):
        safe_name = f"{safe_name}{storage_ext}"
    path = _image_dir() / f"{attachment_id}{storage_ext}"
    path.write_bytes(data)
    return ExtractedAttachment(
        filename=safe_name,
        file_type=storage_ext.lstrip("."),
        text="",
        char_count=0,
        kind="image",
        attachment_id=attachment_id,
        mime_type=detected_mime,
        preview_url=f"/api/inspiration/attachments/{attachment_id}/preview",
    )


def _decode_text(data: bytes) -> str:
    for encoding in ("utf-8-sig", "utf-8", "gb18030", "latin-1"):
        try:
            return data.decode(encoding)
        except UnicodeDecodeError:
            continue
    return data.decode("utf-8", errors="ignore")


def _extract_pdf_text(data: bytes) -> str:
    try:
        from pdfminer.high_level import extract_text
    except Exception as exc:
        raise AttachmentExtractionError("PDF 解析组件未安装，暂时无法读取 PDF。", 500) from exc

    temp_path = None
    try:
        with tempfile.NamedTemporaryFile(delete=False, suffix=".pdf") as handle:
            handle.write(data)
            temp_path = handle.name
        return extract_text(temp_path) or ""
    except Exception as exc:
        raise AttachmentExtractionError("PDF 解析失败，请换一个可复制文字的 PDF。", 400) from exc
    finally:
        if temp_path:
            try:
                os.unlink(temp_path)
            except OSError:
                pass


def _extract_docx_text(data: bytes) -> str:
    try:
        from docx import Document
    except Exception as exc:
        raise AttachmentExtractionError("Word 解析组件未安装，暂时无法读取 Word。", 500) from exc

    try:
        document = Document(BytesIO(data))
    except Exception as exc:
        raise AttachmentExtractionError("Word 文件解析失败，请上传 .docx 文件。", 400) from exc

    parts = [paragraph.text for paragraph in document.paragraphs if paragraph.text.strip()]
    for table in document.tables:
        for row in table.rows:
            cells = [cell.text.strip() for cell in row.cells if cell.text.strip()]
            if cells:
                parts.append(" | ".join(cells))
    return "\n".join(parts)


def _extract_xlsx_text(data: bytes) -> str:
    try:
        import pandas as pd
    except Exception as exc:
        raise AttachmentExtractionError("表格解析组件未安装，暂时无法读取 Excel。", 500) from exc

    try:
        sheets = pd.read_excel(BytesIO(data), sheet_name=None, nrows=80)
    except Exception as exc:
        raise AttachmentExtractionError("Excel 解析失败，请上传 .xlsx 文件。", 400) from exc

    parts = []
    for sheet_name, frame in sheets.items():
        if frame.empty:
            continue
        parts.append(f"【{sheet_name}】")
        parts.append(frame.fillna("").to_csv(index=False))
    return "\n".join(parts)


def extract_attachment_text(filename: str, content_type: str, data: bytes) -> ExtractedAttachment:
    ext = _extension(filename)
    if ext in ALL_IMAGE_EXTENSIONS or _normalized_content_type(content_type).startswith("image/"):
        raise AttachmentExtractionError("图片上传先不做，当前请上传 PDF、Word、文本或表格文件。", 415)
    if ext not in SUPPORTED_EXTENSIONS:
        raise AttachmentExtractionError("暂不支持该文件类型，请上传 PDF、Word、TXT、Markdown、JSON、CSV 或 XLSX。", 415)
    if len(data) > MAX_ATTACHMENT_BYTES:
        raise AttachmentExtractionError("文件过大，请控制在 12MB 以内。", 413)
    if not data:
        raise AttachmentExtractionError("文件为空，请重新选择。", 400)

    if ext in TEXT_EXTENSIONS:
        text = _decode_text(data)
    elif ext == ".pdf":
        text = _extract_pdf_text(data)
    elif ext == ".docx":
        text = _extract_docx_text(data)
    elif ext == ".xlsx":
        text = _extract_xlsx_text(data)
    else:
        text = ""

    text = _trim_text(text)
    if not text:
        raise AttachmentExtractionError("没有从文件中提取到可用文字。", 400)
    return ExtractedAttachment(
        filename=Path(filename or "attachment").name,
        file_type=ext.lstrip("."),
        text=text,
        char_count=len(text),
    )


def extract_inspiration_attachment(filename: str, content_type: str, data: bytes) -> ExtractedAttachment:
    ext = _extension(filename)
    if ext in ALL_IMAGE_EXTENSIONS or _normalized_content_type(content_type).startswith("image/"):
        return _save_image_attachment(filename, content_type, data)
    return extract_attachment_text(filename, content_type, data)


def resolve_image_attachment(attachment_id: str) -> tuple[Path, str]:
    safe_id = _safe_uuid(attachment_id)
    for ext, mime_type in IMAGE_MIME_BY_EXTENSION.items():
        path = _image_dir() / f"{safe_id}{ext}"
        if path.exists() and path.is_file():
            return path, mime_type
    raise AttachmentExtractionError("图片附件已失效，请重新上传。", 400)


def load_image_attachment_data_url(attachment_id: str) -> str:
    path, mime_type = resolve_image_attachment(attachment_id)
    payload = base64.b64encode(path.read_bytes()).decode("ascii")
    return f"data:{mime_type};base64,{payload}"
